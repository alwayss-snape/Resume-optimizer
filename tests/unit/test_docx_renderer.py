import os
import docx
import pytest

from app.analysis.rewriter import RewriteProposal
from app.domain.resume import Candidate, Experience, Resume, ResumeBullet
from app.ingestion.docx import DocxParser
from app.rendering.docx_patcher import DocxPatcher
from app.rendering.template_renderer import TemplateRenderer

def test_docx_patcher_preserve_mode(tmp_path):
    sample_src = "tests/fixtures/resumes/sample.docx"
    parser = DocxParser()
    raw_doc = parser.parse(sample_src)

    first_bullet = [b for b in raw_doc.blocks if b.block_type == "bullet"][0]
    proposal = RewriteProposal(
        source_id=first_bullet.id,
        original_text=first_bullet.text,
        rewritten_text="Architected high-throughput microservices using Python, FastAPI, and Redis.",
        evidence_ids=["ev_1"],
    )

    out_docx = str(tmp_path / "patched.docx")
    patcher = DocxPatcher()
    res_path = patcher.patch(sample_src, raw_doc.document_map, [proposal], out_docx)

    assert os.path.exists(res_path)
    patched_doc = docx.Document(res_path)
    full_text = "\n".join([p.text for p in patched_doc.paragraphs])
    assert "Architected high-throughput microservices using Python, FastAPI, and Redis." in full_text

def test_template_renderer_ats_mode(tmp_path):
    resume = Resume(
        candidate=Candidate(name="John Smith", email="john@example.com"),
        summary="Experienced Software Engineer",
        experience=[
            Experience(
                id="exp_1",
                company="TechCorp",
                title="Lead Developer",
                bullets=[ResumeBullet(id="b1", text="Led cloud architecture on AWS.")],
            )
        ],
    )

    renderer = TemplateRenderer()
    out_path = str(tmp_path / "ats_resume.docx")
    res_path = renderer.render_ats_default(resume, out_path)

    assert os.path.exists(res_path)
    rendered_doc = docx.Document(res_path)
    full_text = "\n".join([p.text for p in rendered_doc.paragraphs])
    assert "John Smith" in full_text
    assert "Led cloud architecture on AWS." in full_text

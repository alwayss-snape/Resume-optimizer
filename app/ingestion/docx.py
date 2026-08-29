import os
import re
from typing import Any, Dict, List, Tuple
import docx
from pydantic import BaseModel, Field

from app.rendering.document_map import DocumentLocation, DocumentMap

class RawBlock(BaseModel):
    id: str
    block_type: str  # 'heading', 'paragraph', 'bullet', 'table_cell'
    text: str
    section: str = "general"
    location: DocumentLocation

class RawDocument(BaseModel):
    filename: str
    blocks: List[RawBlock] = Field(default_factory=list)
    document_map: DocumentMap = Field(default_factory=DocumentMap)
    raw_text: str = ""

class DocxParser:
    BULLET_PREFIXES = ("•", "-", "*", "▪", "–", "—", "o ")

    def parse(self, file_path: str) -> RawDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at path: {file_path}")

        doc = docx.Document(file_path)
        blocks: List[RawBlock] = []
        document_map = DocumentMap()
        full_text_lines: List[str] = []

        current_section = "Header"
        block_counter = 0

        for p_idx, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name if p.style else ""
            is_bullet = (
                style_name.lower().startswith("list bullet")
                or text.startswith(self.BULLET_PREFIXES)
            )
            is_heading = (
                style_name.lower().startswith("heading")
                or (len(text) < 60 and p.runs and any(r.bold for r in p.runs) and not is_bullet)
            )

            if is_heading:
                current_section = text
                block_type = "heading"
            elif is_bullet:
                block_type = "bullet"
                # Strip leading bullet symbol if present
                for prefix in self.BULLET_PREFIXES:
                    if text.startswith(prefix):
                        text = text[len(prefix):].strip()
                        break
            else:
                block_type = "paragraph"

            block_id = f"blk_{block_counter:04d}"
            block_counter += 1

            run_indices = list(range(len(p.runs)))
            location = DocumentLocation(
                section=current_section,
                paragraph_index=p_idx,
                run_indices=run_indices,
                style_name=style_name,
                original_text=p.text,
            )

            raw_block = RawBlock(
                id=block_id,
                block_type=block_type,
                text=text,
                section=current_section,
                location=location,
            )
            blocks.append(raw_block)
            document_map.add_location(block_id, location)
            full_text_lines.append(p.text)

        # Process tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue
                    block_id = f"tbl_{t_idx}_r{r_idx}_c{c_idx}"
                    location = DocumentLocation(
                        section=current_section,
                        table_index=t_idx,
                        row=r_idx,
                        column=c_idx,
                        original_text=cell.text,
                    )
                    raw_block = RawBlock(
                        id=block_id,
                        block_type="table_cell",
                        text=cell_text,
                        section=current_section,
                        location=location,
                    )
                    blocks.append(raw_block)
                    document_map.add_location(block_id, location)
                    full_text_lines.append(cell_text)

        return RawDocument(
            filename=os.path.basename(file_path),
            blocks=blocks,
            document_map=document_map,
            raw_text="\n".join(full_text_lines),
        )

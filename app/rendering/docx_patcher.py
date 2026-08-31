import os
from typing import List

import docx

from app.analysis.rewriter import RewriteProposal
from app.rendering.document_map import DocumentMap

class DocxPatcher:
    """Apply approved text-only rewrites without rebuilding the source document."""

    @staticmethod
    def _replace_paragraph_text(paragraph, text: str) -> None:
        if paragraph.runs:
            first_run = paragraph.runs[0]
            # Retain the leading literal marker for manually-entered bullets. Word
            # list bullets have no marker in paragraph.text and keep their style.
            original = paragraph.text.strip()
            prefix = original[0] + " " if original.startswith(("•", "-", "*", "▪", "–", "—")) else ""
            first_run.text = prefix + text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = text

    def patch(self, source_docx_path: str, document_map: DocumentMap,
              approved_proposals: List[RewriteProposal], output_docx_path: str) -> str:
        if not os.path.exists(source_docx_path):
            raise FileNotFoundError(f"Source DOCX not found at path: {source_docx_path}")

        doc = docx.Document(source_docx_path)
        for proposal in approved_proposals:
            prop_source = getattr(proposal, "source_id", None) or getattr(proposal, "target_source_location_id", None)
            location = document_map.get_location(prop_source)
            if not location:
                # Fail loudly when a proposal targets an unknown source_location_id
                raise ValueError(f"Unknown source location for proposal: {prop_source}")

            if location.paragraph_index is not None:
                if 0 <= location.paragraph_index < len(doc.paragraphs):
                    new_text = getattr(proposal, "rewritten_text", None) or getattr(proposal, "proposed_text", "")
                    self._replace_paragraph_text(doc.paragraphs[location.paragraph_index], new_text)
                continue

            # Résumés often use tables for two-column layouts. Patch the existing
            # cell in place so its widths, borders, fonts, and spacing survive.
            if None not in (location.table_index, location.row, location.column):
                if 0 <= location.table_index < len(doc.tables):
                    table = doc.tables[location.table_index]
                    if 0 <= location.row < len(table.rows) and 0 <= location.column < len(table.rows[location.row].cells):
                        cell = table.rows[location.row].cells[location.column]
                        paragraph_index = location.cell_paragraph_index
                        if paragraph_index is not None and 0 <= paragraph_index < len(cell.paragraphs):
                            new_text = getattr(proposal, "rewritten_text", None) or getattr(proposal, "proposed_text", "")
                            self._replace_paragraph_text(cell.paragraphs[paragraph_index], new_text)

        os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
        doc.save(output_docx_path)
        return output_docx_path

import os

import docx
<<<<<<< HEAD

=======
from typing import Any
>>>>>>> b495747 (Refactor: renderer accepts ResumeDocument; Tailor passes ResumeDocument; add validation agent and tests)
from app.domain.resume import Resume
try:
    from app.domain.resume_document import ResumeDocument
except Exception:
    ResumeDocument = None


class TemplateRenderer:
<<<<<<< HEAD
    """Standard single-column, ATS-safe DOCX renderer with no content omissions."""

    def render_ats_default(self, resume: Resume, output_path: str) -> str:
=======
    def render_ats_default(self, resume_or_doc: Any, output_path: str) -> str:
        # Accept either a Resume or a ResumeDocument; unwrap when necessary
        if hasattr(resume_or_doc, "resume"):
            resume = resume_or_doc.resume
        else:
            resume = resume_or_doc

>>>>>>> b495747 (Refactor: renderer accepts ResumeDocument; Tailor passes ResumeDocument; add validation agent and tests)
        doc = docx.Document()
        doc.add_paragraph(resume.candidate.name, style="Title")

        contact = [
            value for value in (resume.candidate.email, resume.candidate.phone,
                                resume.candidate.location, *resume.candidate.links)
            if value
        ]
        if contact:
            doc.add_paragraph(" | ".join(contact))

<<<<<<< HEAD
        if resume.summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(resume.summary)

        if resume.experience:
=======
        # Contact info
        contact_parts = []
        if resume.candidate.email:
            contact_parts.append(f"Email: {resume.candidate.email}")
        if resume.candidate.phone:
            contact_parts.append(f"Phone: {resume.candidate.phone}")
        if resume.candidate.location:
            contact_parts.append(f"Location: {resume.candidate.location}")
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

        # Summary
        if resume and getattr(resume, "summary", None):
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(resume.summary)

        # Experience
        if resume and getattr(resume, "experience", None):
>>>>>>> b495747 (Refactor: renderer accepts ResumeDocument; Tailor passes ResumeDocument; add validation agent and tests)
            doc.add_heading("Work Experience", level=1)
            for exp in resume.experience:
                heading = " — ".join(value for value in (exp.title, exp.company) if value)
                if exp.start_date or exp.end_date:
                    heading += f" | {' — '.join(value for value in (exp.start_date, exp.end_date) if value)}"
                doc.add_paragraph(heading).runs[0].bold = True
                for bullet in exp.bullets:
                    doc.add_paragraph(bullet.text, style="List Bullet")

        if resume.projects:
            doc.add_heading("Projects", level=1)
            for project in resume.projects:
                doc.add_paragraph(project.name).runs[0].bold = True
                if project.description:
                    doc.add_paragraph(project.description)
                if project.technologies:
                    doc.add_paragraph("Technologies: " + ", ".join(project.technologies))
                for bullet in project.bullets:
                    doc.add_paragraph(bullet.text, style="List Bullet")

<<<<<<< HEAD
        if resume.skills:
            doc.add_heading("Skills", level=1)
            for category, skills in resume.skills.items():
                doc.add_paragraph(f"{category}: {', '.join(skills)}")

        if resume.education:
=======
        # Skills
        if resume and getattr(resume, "skills", None):
            doc.add_heading("Technical Skills", level=1)
            for category, skill_list in resume.skills.items():
                doc.add_paragraph(f"{category}: {', '.join(skill_list)}")

        # Education
        if resume and getattr(resume, "education", None):
>>>>>>> b495747 (Refactor: renderer accepts ResumeDocument; Tailor passes ResumeDocument; add validation agent and tests)
            doc.add_heading("Education", level=1)
            for education in resume.education:
                values = [education.degree, education.institution, education.dates]
                doc.add_paragraph(" — ".join(value for value in values if value))

        if resume.certifications:
            doc.add_heading("Certifications", level=1)
            for certification in resume.certifications:
                doc.add_paragraph(" — ".join(value for value in certification.values() if value))

        if resume.achievements:
            doc.add_heading("Achievements", level=1)
            for achievement in resume.achievements:
                doc.add_paragraph(achievement, style="List Bullet")

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path

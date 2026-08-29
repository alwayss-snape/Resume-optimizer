#!/usr/bin/env python3
"""Generate sample fixture DOCX resume for testing."""

import os
import docx

def create_sample_docx(output_path: str):
    doc = docx.Document()
    
    # Title / Header
    p_head = doc.add_paragraph("Jane Doe")
    p_head.style = 'Title'
    doc.add_paragraph("Email: jane.doe@example.com | Phone: 555-0199 | San Francisco, CA")
    
    # Summary Section
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Senior Software Engineer with 6+ years of experience building scalable backend services and microservices using Python, PostgreSQL, and AWS."
    )
    
    # Experience Section
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Acme Corp — Senior Backend Engineer (2021 – Present)")
    p_b1 = doc.add_paragraph("• Architected high-throughput microservices in Python and FastAPI, handling 50M+ daily API requests.", style='List Bullet')
    p_b2 = doc.add_paragraph("• Reduced database latency by 35% through PostgreSQL index optimization and Redis caching strategy.", style='List Bullet')
    p_b3 = doc.add_paragraph("• Led a team of 4 engineers in migrating monolithic legacy services to Docker and Kubernetes on AWS.", style='List Bullet')
    
    doc.add_paragraph("Beta Systems — Software Engineer (2018 – 2021)")
    p_b4 = doc.add_paragraph("• Built automated ETL data pipelines processing 2TB daily log data using Python, Pandas, and S3.", style='List Bullet')
    p_b5 = doc.add_paragraph("• Implemented CI/CD deployment automation using GitHub Actions and Terraform.", style='List Bullet')

    # Skills Section
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Languages: Python, SQL, JavaScript, Bash")
    doc.add_paragraph("Frameworks & Tools: FastAPI, Django, Docker, Kubernetes, AWS, PostgreSQL, Redis, Git")

    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.S. in Computer Science — University of California, Berkeley (2014 – 2018)")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Sample DOCX created at: {output_path}")

if __name__ == "__main__":
    create_sample_docx("tests/fixtures/resumes/sample.docx")
